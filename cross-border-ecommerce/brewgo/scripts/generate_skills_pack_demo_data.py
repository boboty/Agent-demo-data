#!/usr/bin/env python3
"""Generate deterministic fictional demo inputs for the BrewGo Skills Pack."""
from __future__ import annotations

import csv
import shutil
from datetime import date, timedelta
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

ROOT = Path(__file__).resolve().parents[1] / "classroom" / "skills-pack"


def prepare(skill: str) -> Path:
    demo = ROOT / skill / "demo"
    input_dir = demo / "input"
    outputs = demo / "outputs"
    if input_dir.exists():
        shutil.rmtree(input_dir)
    input_dir.mkdir(parents=True)
    outputs.mkdir(parents=True, exist_ok=True)
    existing = [path for path in outputs.iterdir() if path.name != ".gitkeep"]
    if existing:
        raise SystemExit(f"Refusing to regenerate while demo outputs exist: {outputs}")
    (outputs / ".gitkeep").write_text("", encoding="utf-8")
    return input_dir


def write_text(path: Path, content: str) -> None:
    path.write_text(content.strip() + "\n", encoding="utf-8")


def write_csv(path: Path, headers: list[str], rows: list[list[object]]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(headers)
        writer.writerows(rows)


def write_xlsx(path: Path, sheet: str, headers: list[str], rows: list[list[object]]) -> None:
    workbook = Workbook()
    ws = workbook.active
    ws.title = sheet
    ws.append(headers)
    for row in rows:
        ws.append(row)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="174B35")
        cell.alignment = Alignment(vertical="center")
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    for column, header in enumerate(headers, 1):
        values = [str(header), *(str(ws.cell(row, column).value or "") for row in range(2, ws.max_row + 1))]
        ws.column_dimensions[get_column_letter(column)].width = min(max(len(value) for value in values) + 3, 46)
    workbook.save(path)


def write_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    document = Document()
    document.add_heading(title, 0)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.core_properties.title = title
    document.core_properties.subject = "FICTIONAL CLASSROOM DATA"
    document.save(path)


def write_pdf(path: Path, title: str, lines: list[str]) -> None:
    pdf = canvas.Canvas(str(path), pagesize=A4)
    pdf.setTitle(title)
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(52, 790, title)
    pdf.setFont("Helvetica", 10)
    y = 760
    for line in ["FICTIONAL CLASSROOM DATA", *lines]:
        pdf.drawString(52, y, line[:115])
        y -= 18
    pdf.save()


def write_image(path: Path, heading: str, subheading: str, colour: tuple[int, int, int]) -> None:
    image = Image.new("RGB", (1000, 700), (247, 247, 243))
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((330, 95, 670, 565), radius=45, fill=colour, outline=(28, 35, 31), width=7)
    draw.ellipse((405, 155, 595, 345), fill=(215, 218, 211), outline=(28, 35, 31), width=5)
    draw.rectangle((385, 440, 615, 515), fill=(225, 227, 221), outline=(28, 35, 31), width=5)
    draw.text((45, 40), heading, fill=(20, 27, 23))
    draw.text((45, 650), subheading, fill=(75, 82, 77))
    image.save(path, quality=92)


def generate_reviews() -> None:
    root = prepare("amazon-review-insights")
    rows: list[list[object]] = []
    start = date(2026, 5, 1)
    groups = [
        (12, 5, "Consistent grind", "The grind is even for my pour-over and the cup tastes more consistent than with my old grinder."),
        (10, 5, "Great for travel", "Compact enough for my work bag and camping kit; I like that it needs no battery."),
        (8, 4, "Solid hand feel", "The body feels secure in my hand and the folding handle is comfortable for one cup."),
        (6, 4, "Quick routine clean", "The included brush clears loose grounds quickly after my morning coffee."),
        (8, 3, "Capacity is limited", "Thirty grams is smaller than I expected when making coffee for several people."),
        (6, 2, "Adjustment takes practice", "The adjustment clicks are not obvious at first and I lost my previous setting."),
        (4, 2, "Box arrived damaged", "The shipping box was crushed even though the grinder inside still worked."),
        (3, 2, "Expected an electric grinder", "I thought this would grind automatically, but it is a manual hand grinder."),
        (2, 2, "Grounds in the gap", "Fine grounds collect around the adjustment area and are harder to brush out."),
        (1, 5, "Capacity works for me", "The 30 gram capacity is just right for my two small cups, so I do not see it as too small."),
    ]
    review_number = 1
    variants = ["G2 Black", "G2 Silver", "G2 Travel Kit"]
    for count, rating, title, text in groups:
        for index in range(count):
            rows.append([
                f"R{review_number:03d}", rating, title,
                f"{text} Review sample {index + 1} adds no new product specification.",
                "Y" if review_number % 7 else "N",
                (start + timedelta(days=review_number * 2)).isoformat(),
                variants[(review_number - 1) % len(variants)],
            ])
            review_number += 1
    write_csv(root / "reviews.csv", ["review_id", "rating", "title", "review_text", "verified_purchase", "review_date", "variant"], rows)


def generate_returns() -> None:
    root = prepare("amazon-return-reduction")
    rows: list[list[object]] = []
    start = date(2026, 7, 1)
    categories = [
        (10, "G2-BLK", "Product defective", "Handle became loose after a few uses; requesting return.", 59.99),
        (8, "G2-SLV", "Not as expected", "Capacity is smaller than expected for making coffee for the family.", 59.99),
        (7, "G2-BLK", "Damaged on arrival", "Outer box was crushed and the catch cup arrived dented.", 59.99),
        (6, "G2-BLK", "Difficult to use", "Could not understand how to return to the previous grind setting.", 59.99),
        (5, "G2-TRAVEL", "No longer needed", "Changed my mind after ordering; product was not opened.", 69.99),
        (4, "G2-SLV", "Other", "", 59.99),
    ]
    number = 1
    for count, sku, reason, comment, amount in categories:
        for index in range(count):
            actual_comment = comment
            if reason == "Not as expected" and index >= 5:
                actual_comment = "I expected an electric grinder and did not realise the handle must be turned manually."
            rows.append([
                f"ORD-R{number:04d}", sku, f"B0BG{(number % 3) + 10000}",
                (start + timedelta(days=number)).isoformat(), reason, actual_comment,
                amount, "FBA" if number % 4 else "FBM",
            ])
            number += 1
    write_xlsx(root / "returns.xlsx", "Returns", ["order_id", "sku", "asin", "return_date", "return_reason", "customer_comment", "refund_amount", "fulfillment_channel"], rows)


def generate_inventory() -> None:
    root = prepare("amazon-inventory-watch")
    headers = ["sku", "asin", "current_stock", "reserved_stock", "inbound_qty", "inbound_eta", "avg_daily_sales_7d", "avg_daily_sales_30d", "lead_time_days", "safety_stock_days"]
    rows = [
        ["G2-BLK", "B0BG10001", 82, 12, 0, "", 5.0, 4.0, 25, 14],
        ["G2-SLV", "B0BG10002", 50, 5, 0, "", 2.4, 2.0, 30, 14],
        ["G2-TRAVEL", "B0BG10003", 22, 4, 0, "", 1.7, 1.5, 20, 10],
        ["G2-BAG", "B0BG10004", 1240, 40, 0, "", 3.2, 3.0, 18, 10],
        ["G2-BRUSH", "B0BG10005", 2550, 50, 0, "", 5.5, 5.0, 12, 7],
        ["G2-BOX", "B0BG10006", 40, 10, 500, "", 3.2, 3.0, 14, 10],
        ["G2-HANDLE", "B0BG10007", 210, 20, 300, "2026-09-05", 4.0, 3.8, "", 10],
        ["G2-NEW-RED", "B0BG10008", 120, 5, 0, "", 1.2, "", 28, 14],
        ["G2-BURR", "B0BG10009", 260, 10, 0, "", 8.0, 2.5, 20, 10],
        ["G2-LID", "B0BG10010", 180, 15, 120, "2026-09-01", 3.0, 2.8, 18, 10],
        ["G2-JAR", "B0BG10011", 340, 20, 0, "", 3.4, 3.2, 20, 10],
        ["G2-MANUAL", "B0BG10012", 1600, 100, 0, "", 18.0, 17.0, 10, 5],
        ["G2-SLEEVE", "B0BG10013", 420, 25, 0, "", 5.0, 4.8, 15, 7],
        ["G2-SEAL", "B0BG10014", 300, 40, 200, "2026-08-29", 6.5, 6.0, 12, 7],
        ["G2-CAP", "B0BG10015", 210, 30, 0, "", 4.2, 4.0, 18, 10],
        ["G2-BASE", "B0BG10016", 190, 20, 150, "2026-09-08", 3.1, 3.0, 22, 10],
        ["G2-PIN", "B0BG10017", 900, 40, 0, "", 12.0, 11.5, 14, 7],
        ["G2-CASE", "B0BG10018", 250, 20, 0, "", 3.0, 2.9, 18, 10],
        ["G2-KNOB", "B0BG10019", 380, 30, 0, "", 4.4, 4.2, 20, 10],
        ["G2-GIFT", "B0BG10020", 150, 10, 200, "2026-09-10", 2.8, 2.6, 25, 14],
    ]
    write_xlsx(root / "inventory.xlsx", "Inventory", headers, rows)


def generate_localizer() -> None:
    root = prepare("amazon-listing-localizer")
    write_text(root / "product-facts.md", """
# BrewGo G2 verified product facts

- Product: BrewGo G2 manual burr coffee grinder
- Model: G2
- Operation: manual hand crank; no motor, battery, or charging
- Burr material: stainless steel
- Body: ABS body with stainless steel components
- Bean capacity: 30 g / 1.06 oz
- Product size: 7.8 x 2.4 x 2.4 in / 19.8 x 6.1 x 6.1 cm
- Product weight: 21.9 oz / 620 g
- Included: grinder, folding handle, cleaning brush, travel pouch, user guide
- Supported contexts: home, office, travel, camping; experience may vary
- Warranty note: the supplied one-year limited warranty wording is approved for Amazon US only. UK terms are not supplied.
- No evidence supplied for TSA eligibility, dishwasher safety, waterproofing, titanium burrs, or quantified grind uniformity.
""")
    write_text(root / "listing-us.md", """
# Amazon US Listing — source copy

## Title
BrewGo G2 Manual Coffee Grinder, 1.06 oz Capacity, 7.8 Inch Portable Burr Grinder with Travel Pouch, Black

## Bullets

1. YOUR FAVORITE COFFEE, ANYWHERE — A compact manual grinder for the kitchen, office, camping, and coast-to-coast road trips.
2. 1.06 OZ SINGLE-BATCH CAPACITY — Grind up to 1.06 oz of beans in the 7.8 x 2.4 x 2.4 inch body.
3. NO OUTLET NEEDED — The folding hand crank works without batteries or charging.
4. READY FOR THE ROAD — At 21.9 oz, it packs with the included travel pouch and cleaning brush. TSA-friendly for carry-on travel.
5. US PEACE OF MIND — Covered by a one-year limited warranty for US customers.

## Description
Choose your favorite beans, turn the comfortable handle, and enjoy fresh coffee at home or on the go. The stainless steel burr and compact black body make BrewGo G2 a practical companion for everyday American adventures.
""")


def generate_a_plus() -> None:
    root = prepare("amazon-a-plus-planner")
    write_text(root / "product-profile.md", """
# BrewGo G2 product profile

- Manual hand-crank burr grinder; no battery or charging.
- Stainless steel burr. ABS body with stainless steel components.
- Capacity: 30 g. Weight: 620 g. Size: 19.8 x 6.1 x 6.1 cm.
- Included: folding handle, cleaning brush, travel pouch, user guide.
- Intended contexts supported by product design: home, office, travel, camping.
- Black and silver variants have the same functional specification.
- Do not claim titanium burrs, waterproofing, dishwasher safety, quantified uniformity, or effortless cleaning.
""")
    write_text(root / "listing-current.md", """
# Current Listing excerpt

## Title
BrewGo G2 Portable Coffee Grinder with Premium Titanium Burrs, Easy-Clean Travel Coffee Mill

## Bullets
- Compact 30 g capacity for coffee anywhere.
- Premium titanium burrs create perfectly uniform grounds every time.
- Easy to clean in seconds with the included brush.
- Manual operation means no battery is required.

This copy has not yet been checked against the product profile.
""")
    review_rows = []
    comments = [
        (5, "Fits my office drawer and works well for one cup."),
        (5, "I take it camping because there is no battery to charge."),
        (4, "The folding handle and pouch make packing straightforward."),
        (5, "Grind felt consistent for my pour-over routine."),
        (3, "Thirty grams is not enough when friends visit."),
        (2, "I could not remember the adjustment setting."),
        (4, "The loose grounds brush away quickly."),
        (2, "Fine grounds remain around the adjustment gap."),
        (5, "The black version feels solid in my hand."),
        (4, "Silver looks good on my desk."),
    ]
    for number in range(1, 21):
        rating, comment = comments[(number - 1) % len(comments)]
        review_rows.append([f"AP{number:03d}", rating, comment, "G2 Black" if number % 2 else "G2 Silver"])
    write_csv(root / "reviews.csv", ["review_id", "rating", "review_text", "variant"], review_rows)
    write_text(root / "image-manifest.md", """
# Available image assets

| Filename | Status | What it shows |
| --- | --- | --- |
| g2-hero-white.jpg | available | Black G2 and included items on white background |
| g2-hand-crank-detail.jpg | available | Folding handle and top mechanism close-up |
| g2-office-scene.jpg | available | Grinder beside a mug and pour-over kit on an office desk |
| g2-burr-exploded-view.png | missing | Requested structure view; no approved asset yet |

No lifestyle camping image, comparison chart, certification graphic, or cleaning-step sequence is supplied.
""")


def generate_supplier_quotes() -> None:
    root = prepare("supplier-quote-compare")
    headers = ["SKU", "unit_price", "currency", "MOQ", "lead_time", "packaging_fee", "tooling_fee", "payment_terms", "shipping_term", "valid_until", "analysis_date", "notes"]
    a = [
        ["G2-BLK", 7.65, "USD", 1200, 28, 0.10, 300, "50% deposit, 50% before shipment", "FOB Shenzhen", "2026-08-21", "2026-08-18", "Lowest stated USD unit price; quote expires soon"],
        ["G2-SLV", 7.90, "USD", 1200, 28, 0.10, 300, "50% deposit, 50% before shipment", "FOB Shenzhen", "2026-08-21", "2026-08-18", "G2-BAG not quoted"],
    ]
    b = [
        ["G2-BLK", 7.95, "USD", 500, 45, 0.20, 0, "30% deposit, 70% before shipment", "EXW Dongguan", "2026-09-30", "2026-08-18", "Freight from factory not included"],
        ["G2-SLV", 8.15, "USD", 500, 45, 0.20, 0, "30% deposit, 70% before shipment", "EXW Dongguan", "2026-09-30", "2026-08-18", "Freight from factory not included"],
        ["G2-BAG", 1.20, "USD", 800, 35, 0.05, 0, "30% deposit, 70% before shipment", "EXW Dongguan", "2026-09-30", "2026-08-18", ""],
    ]
    c = [
        ["G2-BLK", 55.50, "CNY", 600, 24, "", 500, "100% before production", "FOB Shenzhen", "2026-09-15", "2026-08-18", "Custom retail box adds CNY 2.20 per unit; not included above"],
        ["G2-SLV", 57.00, "CNY", 600, 24, "", 500, "100% before production", "FOB Shenzhen", "2026-09-15", "2026-08-18", "Custom retail box adds CNY 2.20 per unit; not included above"],
        ["G2-BAG", 8.40, "CNY", 800, 20, 0, 0, "100% before production", "FOB Shenzhen", "2026-09-15", "2026-08-18", "No exchange rate supplied"],
    ]
    write_xlsx(root / "supplier-a.xlsx", "Quote", headers, a)
    write_xlsx(root / "supplier-b.xlsx", "Quote", headers, b)
    write_xlsx(root / "supplier-c.xlsx", "Quote", headers, c)


def generate_customer_service() -> None:
    root = prepare("customer-service-triage")
    messages = [
        ("01-tracking.eml", "Maya Stone", "Where is order BG-8101? The tracking page has not updated since Monday."),
        ("02-delay.eml", "Owen Reed", "Order BG-8102 was promised last Friday and is still not here. Please check the carrier delay."),
        ("03-use-question.eml", "Leah Park", "How do I return to my previous grind setting after cleaning the G2?"),
        ("04-quality.eml", "Noah King", "The folding handle feels loose after three uses. I can send a video if needed."),
        ("05-refund.eml", "Ava Brooks", "Please approve a full USD 59.99 refund for order BG-8105. The item is unopened."),
        ("06-replacement.eml", "Ethan Bell", "The travel pouch zip is broken. Can you tell me the replacement process for BG-8106?"),
        ("07-review-threat.eml", "Mia Cole", "If nobody responds today I will post a one-star review and escalate this complaint to Amazon."),
        ("08-compensation.eml", "Liam Shaw", "The delayed gift caused embarrassment. I demand USD 500 compensation plus a refund."),
        ("09-safety.eml", "Sophia Lane", "I cut my finger on a sharp edge inside the catch cup. I stopped using it and kept the product."),
        ("10-missing-info.eml", "Guest Customer", "It does not work. Fix this now. I do not have the order number with me."),
        ("11-address.eml", "Lucas Gray", "Can I change the delivery address for BG-8111 before it ships?"),
        ("12-damaged.eml", "Emma Wood", "The retail box and catch cup arrived crushed. Photos are attached to my original message."),
        ("13-legal.eml", "James Fox", "My lawyer says your warranty wording may be misleading. Confirm in writing that BrewGo accepts liability."),
        ("14-use-cleaning.eml", "Isla Moore", "Is the G2 dishwasher safe? I cannot find that statement in the manual."),
    ]
    for filename, customer, body in messages:
        number = filename.split("-", 1)[0]
        write_text(root / filename, f"""
From: {customer} <customer{number}@example.test>
To: BrewGo Support <support@brewgo.example>
Date: 2026-08-{int(number) + 3:02d} 10:15
Subject: Customer request {number}

{body}
""")


def generate_file_organizer() -> None:
    root = prepare("business-file-organizer")
    write_text(root / "Listing_US_final.md", "# BrewGo G2 US Listing draft\nStatus: approved copy draft, not published.\nModel: G2-BLK")
    write_docx(root / "Listing最终2.docx", "BrewGo G2 UK Listing draft", ["Version: 2", "Status: localisation review pending", "Product: G2-BLK manual grinder"])
    write_image(root / "IMG_3381.jpg", "BrewGo G2 sample", "Black sample / angle view / approval pending", (44, 48, 46))
    write_image(root / "product_white_final.png", "BrewGo G2 white background", "Listing image candidate / 2026-08-12", (36, 39, 37))
    headers = ["Supplier", "SKU", "Currency", "Unit Price", "MOQ", "Date"]
    write_xlsx(root / "最终版.xlsx", "Quote", headers, [["BrewPeak", "G2-BLK", "USD", 7.80, 500, "2026-08-05"]])
    write_xlsx(root / "报价最终_final.xlsx", "Quote", headers, [["Aurora", "G2-BLK", "USD", 8.20, 500, "2026-08-16"]])
    write_pdf(root / "new2.pdf", "Commercial Invoice Draft", ["Invoice: INV-BG-0815-DRAFT", "Customer: NorthPeak Coffee", "Amount: USD 7,200", "Status: DRAFT"])
    write_pdf(root / "invoice_0815.pdf", "Commercial Invoice", ["Invoice: INV-BG-0815", "Customer: NorthPeak Coffee", "Amount: USD 7,200", "Status: ISSUED"])
    write_xlsx(root / "PackingList copy.xlsx", "Packing List", ["Shipment", "Carton", "SKU", "Units"], [["BG260815", 1, "G2-BLK", 24], ["BG260815", 2, "G2-BLK", 24]])
    write_csv(root / "reviews_new.csv", ["review_id", "rating", "review_text"], [["R901", 5, "Compact for office use"], ["R902", 2, "Box arrived crushed"]])
    write_xlsx(root / "广告报表final.xlsx", "Ads", ["Campaign", "Clicks", "Spend", "Orders"], [["G2 Manual Exact", 120, 84.50, 9]])
    write_docx(root / "会议记录.docx", "2026-08-16 供应链会议", ["G2-BLK 新报价待审批。", "BG260815 报关资料已完成。", "包装破损需复核承运商记录。"])
    write_docx(root / "客户资料NEW.docx", "NorthPeak Coffee 客户资料", ["Market: US", "Contact: Maya Stone", "Interest: 120-unit corporate gift order", "Address pending confirmation"])
    write_text(root / "product-spec.txt", "Product: BrewGo G2\nModel: G2-BLK\nCapacity: 30 g\nOperation: manual hand crank\nBurr: stainless steel")
    write_image(root / "screenshot_001.png", "Amazon Ads screenshot mock", "Campaign G2 Manual Exact / temporary reference", (88, 111, 92))
    write_pdf(root / "合同草稿.pdf", "Supplier Framework Agreement", ["Parties: BrewGo and BrewPeak", "Status: DRAFT", "Payment and inspection clauses require legal review"])
    write_csv(root / "shipping-update.csv", ["shipment", "status", "updated"], [["BG260815", "Departed port", "2026-08-16"]])
    write_text(root / "note.txt", "0818\nPlease check this before the call.\nNo customer, project, document type, or owner is identified.")
    write_pdf(root / "附件.pdf", "Attachment", ["For review", "Lily", "0817", "No project or business category supplied"])
    (root / "unknown_02.dat").write_bytes(b"FICTIONAL CLASSROOM DATA\nref=BG-X2\nvalue=17\ncontext=not supplied\n")


def main() -> None:
    generate_reviews()
    generate_returns()
    generate_inventory()
    generate_localizer()
    generate_a_plus()
    generate_supplier_quotes()
    generate_customer_service()
    generate_file_organizer()
    print("BrewGo Skills Pack demo inputs generated")


if __name__ == "__main__":
    main()
